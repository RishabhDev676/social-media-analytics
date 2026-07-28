from docx import Document
from docx.shared import Pt, RGBColor
import os

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.style = 'Intense Quote'

def main():
    doc = Document()
    
    # Title
    add_heading(doc, 'Social Media Sentiment Analysis: Full Project Guide', level=0)
    add_paragraph(doc, 'This document explains every file in our project in simple, easy-to-understand terms. This will help you understand how the code works from start to finish.', bold=True)
    
    # Overview
    add_heading(doc, '1. Project Overview', level=1)
    doc.add_paragraph("Our project takes social media comments and predicts whether they are Positive, Negative, or Neutral. We do this in several steps, separated into different Python files:")
    doc.add_paragraph("• preprocessing.py: Cleans up messy text.")
    doc.add_paragraph("• train_model.py: Teaches the AI using our 11,000 comment dataset.")
    doc.add_paragraph("• predict.py: Uses the trained AI to analyze new comments.")
    doc.add_paragraph("• statistics.py: Crunches the numbers to show overall sentiment.")
    doc.add_paragraph("• gui.py: Provides a visual app for the user to upload files.")

    # Imports Explanation
    add_heading(doc, '2. Understanding the Python Imports', level=1)
    add_paragraph(doc, "Across our project, we import several libraries. Here is what each one does:")
    doc.add_paragraph("• pandas: Used to load, manipulate, and analyze our data using DataFrames (like Excel tables in code).")
    doc.add_paragraph("• os: A built-in library to handle file paths and directories, ensuring our code works on Mac, Windows, and Linux.")
    doc.add_paragraph("• re: The Regular Expression library, used for searching and replacing specific text patterns (like removing URLs).")
    doc.add_paragraph("• nltk (Natural Language Toolkit): We use this to download and remove 'stopwords' (common words like 'the' or 'is' that don't add meaning).")
    doc.add_paragraph("• sklearn (Scikit-Learn): Our main Machine Learning library. We import 'TfidfVectorizer' to convert text to numbers, 'LogisticRegression' for the AI model, 'train_test_split' to divide our data, and metrics to check accuracy.")
    doc.add_paragraph("• joblib: Used to save our trained model and vectorizer to disk so we can load them later without retraining.")
    doc.add_paragraph("• tkinter: Python's standard GUI library. We use it to build the desktop window, buttons, and file dialogs.")
    doc.add_paragraph("• numpy: Used for fast mathematical and numerical operations on arrays.")
    doc.add_paragraph("• docx: The library used to generate this Word document programmatically!")

    # Preprocessing
    add_heading(doc, '3. preprocessing.py (Cleaning the Data)', level=1)
    add_paragraph(doc, "People type messily on social media. Before the AI can read it, we have to clean the text.")
    add_code(doc, "text = re.sub(r\"http\S+|www\S+|https\S+\", \"\", text)\ntext = re.sub(r\"[^a-zA-Z\s]\", \"\", text)\nwords = [word for word in words if word not in stop_words]")
    doc.add_paragraph("• We use 're' (Regular Expressions) to search and destroy web links (URLs).")
    doc.add_paragraph("• We delete all punctuation, numbers, and emojis so only English letters remain.")
    doc.add_paragraph("• We use the 'NLTK' library to remove \"stop words.\" Stop words are common words like \"the\", \"and\", or \"is\" that don't tell us anything about sentiment.")

    # Training
    # Training
    add_heading(doc, '4. train_model.py (Teaching the AI)', level=1)
    add_paragraph(doc, "After preprocessing, we move to training. Here is a brief theory of what happens next:", bold=True)
    add_paragraph(doc, "Theory: Machine learning models only understand numbers, not text. Therefore, we must convert our cleaned text into a numerical format. This step is called 'Feature Extraction'. We use TF-IDF (Term Frequency-Inverse Document Frequency), which gives higher importance to unique words that define a sentiment (like 'terrible' or 'amazing') and lowers the importance of common words.")
    add_paragraph(doc, "Once converted to numbers, we split our data: 80% to train the model, and 20% to test it. We use Logistic Regression, which is a statistical model that finds the relationship between the words (features) and the sentiment (labels). It calculates the probability that a comment belongs to a specific category.")
    
    add_paragraph(doc, "Here is the detailed code that accomplishes this:", bold=True)
    
    detailed_train_code = """
# 1. Create the TF-IDF Vectorizer
# ngram_range=(1, 2) means we look at single words and 2-word phrases
# max_features=5000 limits the vocabulary to the top 5000 words
vectorizer = TfidfVectorizer(ngram_range=(1, 2), stop_words='english', max_features=5000)

# 2. Convert text to numbers
# fit_transform learns the vocabulary and transforms the text into a matrix (X)
X = vectorizer.fit_transform(df["cleaned_comment"])

# 3. Define the labels (what we want to predict)
y = df["sentiment"]

# 4. Split data into training (80%) and testing (20%)
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42
)

# 5. Create and Train the Logistic Regression Model
# max_iter=1000 ensures it has enough time to find the best mathematical pattern
model = LogisticRegression(max_iter=1000)
model.fit(X_train, y_train)

# 6. Evaluate and Save
# We predict on the test set to see how well it learned
y_pred = model.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)

# Finally, save the trained model and vectorizer for future use
joblib.dump(model, 'model.pkl')
joblib.dump(vectorizer, 'vectorizer.pkl')
"""
    add_code(doc, detailed_train_code.strip())

    # Prediction
    add_heading(doc, '5. predict.py (Making New Guesses)', level=1)
    add_paragraph(doc, "Now that the AI is smart and saved to a file, this script loads it up to analyze brand-new comments.")
    add_code(doc, "model = joblib.load(model_path)\nvectorizer = joblib.load(vectorizer_path)\n\ntext_features = vectorizer.transform([cleaned_text])\nprediction = model.predict(text_features)")
    doc.add_paragraph("• joblib.load: Wakes up our saved model from the hard drive.")
    doc.add_paragraph("• vectorizer.transform: Turns the new user comment into math numbers exactly the same way we did during training.")
    doc.add_paragraph("• model.predict: The AI looks at the numbers and spits out a final guess: Positive, Negative, or Neutral.")
    doc.add_paragraph("• CSV Export: Finally, it saves these exact predictions into 'predicted_sentiments.csv' so our stats and graphs modules can use them!")

    # Statistics
    add_heading(doc, '6. statistics.py (Crunching the Numbers)', level=1)
    add_paragraph(doc, "Once we have our 'predicted_sentiments.csv', we want to see the big picture.")
    add_code(doc, "counts = df[\"sentiment\"].value_counts()\npositive_pct = (positive_count / total) * 100")
    doc.add_paragraph("• We use the 'Pandas' library to read the predicted CSV and count how many comments fell into each bucket.")
    doc.add_paragraph("• We calculate percentages to show the user exactly what portion of their audience is happy versus unhappy.")

    # GUI
    add_heading(doc, '7. gui.py (The Visual App)', level=1)
    add_paragraph(doc, "Users don't want to type code; they want buttons to click. We built a desktop app using 'Tkinter'.")
    add_code(doc, "import tkinter as tk\nfrom tkinter import filedialog\n\nfile_path = filedialog.askopenfilename()")
    doc.add_paragraph("• Tkinter (tk): This is Python's built-in tool for drawing windows, buttons, and text boxes.")
    doc.add_paragraph("• filedialog: This opens up your computer's standard file browser so the user can visually select their CSV file to upload.")
    doc.add_paragraph("• When the user clicks 'Upload', our app grabs the file, runs it through predict.py, and displays the results!")

    # Save Document
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.join(base_dir, "..", "docs")
    os.makedirs(docs_dir, exist_ok=True)
    
    output_path = os.path.join(docs_dir, "Model_Explanation.docx")
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    main()
