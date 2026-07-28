import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 0:
            run.font.color.rgb = RGBColor(16, 44, 87) # Dark Navy
            run.font.size = Pt(24)
            run.font.bold = True
        elif level == 1:
            run.font.color.rgb = RGBColor(30, 86, 160) # Medium Blue
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.color.rgb = RGBColor(53, 162, 159) # Teal Accent
            run.font.size = Pt(13)
            run.font.bold = True
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.5)
    return p

def add_bullet(doc, title, desc):
    p = doc.add_paragraph(style='List Bullet')
    run_t = p.add_run(title + ": ")
    run_t.bold = True
    run_t.font.name = 'Segoe UI'
    run_t.font.size = Pt(10)
    
    run_d = p.add_run(desc)
    run_d.font.name = 'Segoe UI'
    run_d.font.size = Pt(10)

def main():
    doc = Document()

    # Document Title
    add_heading(doc, 'Social Media Sentiment Analysis: Technical Backlog & Pending Roadmap', level=0)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_m = p_meta.add_run("Target Release: v1.1.0 & v2.0.0\nDocument Type: Open Technical Roadmap\nLast Updated: 2026-07-29")
    run_m.italic = True
    run_m.font.size = Pt(9.5)
    run_m.font.color.rgb = RGBColor(100, 100, 100)

    # Overview
    add_heading(doc, '1. Project Overview & Enhancement Goals', level=1)
    add_paragraph(doc, "This document outlines pending engineering tasks, feature enhancements, optimization roadmaps, and technical debt items for the Social Media Sentiment Analysis project.", bold=False)

    # Task Table
    add_heading(doc, '2. Pending Feature Matrix', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_titles = ['Priority', 'Task / Feature Area', 'Target Package / File', 'Technical Description']
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = shading.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        if shd is None:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1E56A0')
            shading.append(shd)
        else:
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '1E56A0')

    tasks_data = [
        ("High Priority", "Data & Report Export", "src/ui/ui_manager.py\nsrc/utils/helpers.py", "Add one-click export buttons to save analyzed dataset predictions and summary statistics to Excel (.xlsx) and CSV."),
        ("High Priority", "Instant Single-Comment Predictor", "src/ui/ui_manager.py\nsrc/core/predict.py", "Integrate a real-time single comment text box on the home screen for instant sentiment prediction without requiring file uploads."),
        ("Medium Priority", "Multi-Model Benchmarking Suite", "src/core/benchmark.py", "Develop automated evaluation script comparing Logistic Regression against Random Forest, Naive Bayes, and Support Vector Machines (SVM)."),
        ("Medium Priority", "Standalone .exe Distribution Build", "build_exe.py", "Configure PyInstaller packaging script to bundle Python runtime, Tkinter GUI, and trained model binaries into a standalone Windows executable."),
        ("Low Priority", "Multi-Language Processing Support", "src/core/preprocess.py", "Extend text normalization and vectorization pipeline to handle multi-lingual social media text (e.g. Hindi, Spanish).")
    ]

    for row_data in tasks_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[i].paragraphs[0].runs[0].font.name = 'Segoe UI'

    doc.add_paragraph()

    # Detailed Section
    add_heading(doc, '3. Detailed Specifications', level=1)
    
    add_heading(doc, '3.1 One-Click Data & Report Export (v1.1.0)', level=2)
    add_bullet(doc, "Objective", "Allow users to save analyzed results (Original Comment, Cleaned Comment, Predicted Sentiment) to an Excel file (.xlsx).")
    add_bullet(doc, "Implementation Details", "Integrate openpyxl/pandas ExcelWriter in src/utils/helpers.py and attach an 'Export Results' button in SentimentApp.render_dashboard().")

    add_heading(doc, '3.2 Instant Single-Comment Predictor Card (v1.1.0)', level=2)
    add_bullet(doc, "Objective", "Allow users to test individual sentences instantly on the dashboard landing view.")
    add_bullet(doc, "Implementation Details", "Connect Tkinter Text entry widget directly to predict_single_comment() in src/core/predict.py and render a real-time prediction card.")

    add_heading(doc, '3.3 Standalone Executable Packaging (v1.2.0)', level=2)
    add_bullet(doc, "Objective", "Package the entire application into a standalone Windows executable (.exe) for distribution.")
    add_bullet(doc, "Implementation Details", "Use PyInstaller with --noconsole --onefile flags and ensure relative path resolution for sentiment_model.pkl and vectorizer.pkl.")

    # Save
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.join(base_dir, "..", "..", "docs")
    os.makedirs(docs_dir, exist_ok=True)
    output_path = os.path.join(docs_dir, "Pending_Work.docx")
    doc.save(output_path)
    print(f"Pending_Work.docx successfully created at: {output_path}")

if __name__ == "__main__":
    main()
