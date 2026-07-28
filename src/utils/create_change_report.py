import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 0:
            run.font.color.rgb = RGBColor(16, 44, 87) # Dark Navy
            run.font.size = Pt(24)
            run.font.bold = True
        elif level == 1:
            run.font.color.rgb = RGBColor(30, 86, 160) # Medium Blue
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.color.rgb = RGBColor(53, 162, 159) # Teal Accent
            run.font.size = Pt(13)
            run.font.bold = True
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.5)
    return p

def add_bullet(doc, title, desc):
    p = doc.add_paragraph(style='List Bullet')
    run_t = p.add_run(title + ": ")
    run_t.bold = True
    run_t.font.name = 'Segoe UI'
    run_t.font.size = Pt(10)
    
    run_d = p.add_run(desc)
    run_d.font.name = 'Segoe UI'
    run_d.font.size = Pt(10)

def main():
    doc = Document()

    # Document Header Title
    add_heading(doc, 'Social Media Sentiment Analysis: Master Progress Report & Team Guide Index', level=0)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_m = p_meta.add_run("Team Members: Rishabh, Swarni, Daksha, Prathamesh, Manvi, Sai\nTimeline: Day 1 (2026-07-22) to Tomorrow (2026-07-30)\nLast Updated: 2026-07-29")
    run_m.italic = True
    run_m.font.size = Pt(9.5)
    run_m.font.color.rgb = RGBColor(100, 100, 100)

    # Executive Overview
    add_heading(doc, '1. Executive Overview & Index', level=1)
    add_paragraph(doc, "This document provides a comprehensive chronological index of all development phases, team member task allocations from the Team Guide (Pending_Work_Team_Guide.docx), machine learning model enhancements, UI iterations, and future roadmaps for the Social Media Sentiment Analysis project.", bold=False)

    # Report Index Table
    add_heading(doc, 'Report Index Table', level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_titles = ['Day / Date', 'Phase / Focus Area', 'Key Achievements & Team Allocations', 'Status']
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = shading.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        if shd is None:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1E56A0')
            shading.append(shd)
        else:
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '1E56A0')

    reports_data = [
        ("Day 1 (2026-07-22)", "Project Initialization", "Repo setup, README framework, requirements.txt (Rishabh)", "Completed"),
        ("Day 2 (2026-07-26)", "Core ML Pipeline & Team Guide", "Preprocessing, training, predict scripts (Rishabh/Swarni), team guide creation", "Completed"),
        ("Day 3 (2026-07-27)", "GUI Layout & Statistics", "statistics.py (Manvi), GUI layout (Daksha), backend wiring (Prathamesh)", "Completed"),
        ("Day 4 (2026-07-28)", "Model Upgrade & 11k Dataset", "Logistic Regression + TF-IDF, 11k dataset, Model Explanation doc (Rishabh)", "Completed"),
        ("Day 5 Today (2026-07-29)", "Modular Architecture & UI", "Clean package layout, Tkinter SentimentApp dashboard, file/git cleanup", "Completed"),
        ("Day 6 Tomorrow (2026-07-30+)", "Production & Standalone .exe", "Export functions, single-comment live predictor, PyInstaller .exe build", "Planned")
    ]

    for row_data in reports_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[i].paragraphs[0].runs[0].font.name = 'Segoe UI'

    doc.add_paragraph() # Spacer

    # Team Member Work Allocation Section (Extracted from Pending_Work_Team_Guide.docx)
    add_heading(doc, '2. Team Member Work Breakdown (Team Guide)', level=1)
    add_paragraph(doc, "Per the project architecture guide (Pending_Work_Team_Guide.docx), the system modules were distributed across team members as follows:")

    team_table = doc.add_table(rows=1, cols=4)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_hdr = team_table.rows[0].cells
    t_titles = ['Team Member', 'Assigned Module / File', 'Responsibilities & Functions', 'Module Status']
    for i, title in enumerate(t_titles):
        t_hdr[i].text = title
        t_hdr[i].paragraphs[0].runs[0].font.bold = True
        t_hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = t_hdr[i]._tc.get_or_add_tcPr()
        shd = shading.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        if shd is None:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '35A29F')
            shading.append(shd)
        else:
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '35A29F')

    team_rows = [
        ("Rishabh & Swarni", "src/core/preprocess.py\nsrc/core/train.py\nsrc/core/predict.py", "Built regex text cleaner, stop-word filtering, TF-IDF vectorizer, Logistic Regression model trainer, and inference functions.", "Completed"),
        ("Manvi", "src/analytics/statistics.py", "Built compute_statistics() calculating sentiment counts (pos/neg/neu), percentages, text length stats, and frequency distribution tables.", "Completed"),
        ("Sai", "src/analytics/visualization.py", "Built Matplotlib chart generators (create_pie_chart, create_bar_chart, create_histogram) for visual analysis.", "Completed (Optional/Independent)"),
        ("Daksha", "src/ui/ui_manager.py\n(legacy src/gui.py)", "Designed desktop GUI window layout, frame navigation (Home, Results, Graph views), title headers, cards, and buttons.", "Completed"),
        ("Prathamesh", "src/ui/ui_manager.py\n(backend wiring)", "Wired GUI buttons to predict.py, statistics.py, and CSV file upload handling.", "Completed")
    ]

    for row_data in team_rows:
        row_cells = team_table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[i].paragraphs[0].runs[0].font.name = 'Segoe UI'

    doc.add_paragraph() # Spacer

    # Day 1
    add_heading(doc, '3. Day 1 (2026-07-22): Project Initialization', level=1)
    add_paragraph(doc, "On Day 1, the foundational repository environment was established to support machine learning and data processing workflows.")
    add_bullet(doc, "Git Repository Setup", "Created initial repository structure and initialized git version control.")
    add_bullet(doc, "Environment & Dependencies", "Configured requirements.txt defining essential libraries: pandas, scikit-learn, nltk, joblib, matplotlib.")
    add_bullet(doc, "Initial Documentation", "Drafted base README outlining project objectives and planned architecture.")

    # Day 2
    add_heading(doc, '4. Day 2 (2026-07-26): Core ML Pipeline & Team Guide', level=1)
    add_paragraph(doc, "Day 2 focused on constructing the underlying NLP data pipeline and initial batch test data.")
    add_bullet(doc, "Preprocessing Module", "Built text normalization functions using regex for URL and punctuation stripping, plus NLTK stop-word removal.")
    add_bullet(doc, "Model Training & Inference", "Developed standalone training script (train_model.py) and inference script (predict.py).")
    add_bullet(doc, "Test Dataset Creation", "Generated test_100_comments.csv containing 100 sample social media comments for verification.")
    add_bullet(doc, "Team Work Guide Creation", "Created Pending_Work_Team_Guide.docx detailing sub-task breakdowns for team contributors (Manvi, Sai, Daksha, Prathamesh).")

    # Day 3
    add_heading(doc, '5. Day 3 (2026-07-27): GUI Integration & Collaboration', level=1)
    add_paragraph(doc, "Day 3 brought initial graphical user interface development and team code integration.")
    add_bullet(doc, "Statistics Module (Manvi)", "Integrated statistics.py computing positive, negative, and neutral percentages and frequencies.")
    add_bullet(doc, "Visualization Module (Sai)", "Created Matplotlib visualization.py module with pie chart, bar chart, and length histogram functions.")
    add_bullet(doc, "GUI Prototype (Daksha & Prathamesh)", "Integrated GUI layout and button handlers for CSV uploading and statistical displays.")

    # Day 4
    add_heading(doc, '6. Day 4 (2026-07-28): Model Upgrade & 11,000 Comment Dataset', level=1)
    add_paragraph(doc, "Day 4 represented a major machine learning leap forward in accuracy and dataset scale.")
    add_bullet(doc, "Logistic Regression Upgrade", "Replaced baseline model with Logistic Regression combined with TF-IDF Vectorization (unigrams + bigrams, top 5,000 features).")
    add_bullet(doc, "11,000 Comment Dataset", "Generated realistic 11,000-comment dataset (comments.csv) with controlled 5% label noise to emulate real social media data (~95% accuracy).")
    add_bullet(doc, "Batch Inference Export", "Executed batch predictions across the 11,000 comments dataset saving output to predicted_sentiments.csv.")
    add_bullet(doc, "Comprehensive Theory Docs", "Authored Model_Explanation.docx explaining TF-IDF math, Logistic Regression mechanics, and code functions.")

    # Day 5 (Today)
    add_heading(doc, '7. Day 5 Today (2026-07-29): Modular Architecture & Modern UI Redesign', level=1)
    add_paragraph(doc, "Today, the project underwent a complete production-grade refactoring and UI modernization.")
    add_bullet(doc, "Package Restructuring", "Organized project into clean Python packages: src/core (pipeline, model manager, train, predict, preprocess), src/analytics (statistics), src/ui (ui_manager), src/utils (helpers, doc generators).")
    add_bullet(doc, "Desktop Dashboard (SentimentApp)", "Built state-of-the-art Tkinter desktop dashboard in app.py & ui_manager.py featuring real-time positive/negative/neutral count cards (4414 Positive, 4385 Negative, 2201 Neutral), interactive processed data treeviews, and non-blocking background threading.")
    add_bullet(doc, "Clean & Focused UI", "Removed graph visualization clutter from UI and pipeline per user specifications to ensure maximum responsiveness and visual clarity.")
    add_bullet(doc, "Unnecessary File Removal", "Purged redundant scripts (gui.py, visualization.py), MS Word lock files, empty directories, and bytecode caches.")
    add_bullet(doc, "Git History Consolidation", "Squashed trial/back-and-forth commits into a clean, professional commit history and updated GitHub main.")

    # Day 6 (Tomorrow)
    add_heading(doc, '8. Day 6 Tomorrow (2026-07-30+): Production Distribution & Future Roadmap', level=1)
    add_paragraph(doc, "Planned engineering roadmap for tomorrow and upcoming release cycles:")
    add_bullet(doc, "Export Module", "Implement one-click export buttons on the dashboard to save analyzed data and summary tables to Excel (.xlsx) and PDF format.")
    add_bullet(doc, "Single-Comment Live Predictor Card", "Add an interactive instant sentiment scoring card on the dashboard home view for real-time text analysis without uploading files.")
    add_bullet(doc, "Model Comparison Benchmarking", "Add benchmarking scripts in src/core/ to compare Logistic Regression performance against Random Forest, Naive Bayes, and Linear SVM.")
    add_bullet(doc, "Standalone Executable Build", "Package the application using PyInstaller into a standalone Windows executable (.exe) for distribution to end-users without requiring Python installation.")

    # Save
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.join(base_dir, "..", "..", "docs")
    os.makedirs(docs_dir, exist_ok=True)
    output_path = os.path.join(docs_dir, "Change_Report.docx")
    doc.save(output_path)
    print(f"Change_Report.docx successfully created at: {output_path}")

if __name__ == "__main__":
    main()
