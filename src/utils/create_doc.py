from docx import Document
from docx.shared import Pt, RGBColor
import os

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.style = 'Intense Quote'

def main():
    doc = Document()
    
    # Title
    add_heading(doc, 'Social Media Sentiment Analysis: Full Project Guide', level=0)
    add_paragraph(doc, 'This document explains every file in our project in simple, easy-to-understand terms. This will help you understand how the code works from start to finish.', bold=True)
    
    # Overview
    add_heading(doc, '1. Project Overview', level=1)
    doc.add_paragraph("Our project takes social media comments and predicts whether they are Positive, Negative, or Neutral. We do this in several steps, separated into modular packages:")
    doc.add_paragraph("• app.py: Main application entry point.")
    doc.add_paragraph("• src/core/preprocess.py: Cleans up messy text.")
    doc.add_paragraph("• src/core/train.py: Teaches the AI using our 11,000 comment dataset.")
    doc.add_paragraph("• src/core/predict.py: Uses the trained AI to analyze new comments.")
    doc.add_paragraph("• src/core/model_manager.py: Manages loading and automatic training of ML models.")
    doc.add_paragraph("• src/core/pipeline.py: End-to-end pipeline linking preprocessing, prediction, and analytics.")
    doc.add_paragraph("• src/analytics/statistics.py: Crunches the numbers to show overall sentiment and text stats.")
    doc.add_paragraph("• src/ui/ui_manager.py: Tkinter dashboard UI for dataset upload, manual comment entry, and analytics.")

    # Imports Explanation
    add_heading(doc, '2. Understanding the Python Imports', level=1)
    add_paragraph(doc, "Across our project, we import several libraries. Here is what each one does:")
    doc.add_paragraph("• pandas: Used to load, manipulate, and analyze our data using DataFrames (like Excel tables in code).")
    doc.add_paragraph("• os: A built-in library to handle file paths and directories, ensuring our code works on Mac, Windows, and Linux.")
    doc.add_paragraph("• re: The Regular Expression library, used for searching and replacing specific text patterns (like removing URLs).")
    doc.add_paragraph("• nltk (Natural Language Toolkit): We use this to download and remove 'stopwords' (common words like 'the' or 'is' that don't add meaning).")
    doc.add_paragraph("• sklearn (Scikit-Learn): Our main Machine Learning library. We import 'TfidfVectorizer' to convert text to numbers, 'LogisticRegression' for the AI model, 'train_test_split' to divide our data, and metrics to check accuracy.")
    doc.add_paragraph("• joblib: Used to save our trained model and vectorizer to disk so we can load them later without retraining.")
    doc.add_paragraph("• tkinter: Python's standard GUI library. We use it to build the desktop window, buttons, and analytics tables.")
    doc.add_paragraph("• numpy: Used for fast mathematical and numerical operations on arrays.")
    doc.add_paragraph("• docx: The library used to generate this Word document programmatically!")

    # Preprocessing
    add_heading(doc, '3. src/core/preprocess.py (Cleaning the Data)', level=1)
    add_paragraph(doc, "People type messily on social media. Before the AI can read it, we have to clean the text.")
    add_code(doc, "text = re.sub(r\"http\S+|www\S+|https\S+\", \"\", text)\ntext = re.sub(r\"[^a-zA-Z\s]\", \"\", text)\nwords = [word for word in words if word not in stop_words]")
    doc.add_paragraph("• We use 're' (Regular Expressions) to search and destroy web links (URLs).")
    doc.add_paragraph("• We delete all punctuation, numbers, and emojis so only English letters remain.")
    doc.add_paragraph("• We use the 'NLTK' library to remove \"stop words.\" Stop words are common words like \"the\", \"and\", or \"is\" that don't tell us anything about sentiment.")

    # Training
    add_heading(doc, '4. src/core/train.py (Teaching the AI)', level=1)
    add_paragraph(doc, "After preprocessing, we move to training. Here is a brief theory of what happens next:", bold=True)
    add_paragraph(doc, "Theory: Machine learning models only understand numbers, not text. Therefore, we must convert our cleaned text into a numerical format. This step is called 'Feature Extraction'. We use TF-IDF (Term Frequency-Inverse Document Frequency), which gives higher importance to unique words that define a sentiment (like 'terrible' or 'amazing') and lowers the importance of common words.")
    add_paragraph(doc, "Once converted to numbers, we split our data: 80% to train the model, and 20% to test it. We use Logistic Regression, which is a statistical model that finds the relationship between the words (features) and the sentiment (labels). It calculates the probability that a comment belongs to a specific category.")
    
    add_paragraph(doc, "Here is the detailed code that accomplishes this:", bold=True)
    
    detailed_train_code = """
# 1. Create the TF-IDF Vectorizer
vectorizer = TfidfVectorizer(ngram_range=(1, 2), stop_words='english', max_features=5000)

# 2. Convert text to numbers
X = vectorizer.fit_transform(df["cleaned_comment"])

# 3. Define the labels (what we want to predict)
y = df["sentiment"]

# 4. Split data into training (80%) and testing (20%)
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42
)

# 5. Create and Train the Logistic Regression Model
model = LogisticRegression(max_iter=1000)
model.fit(X_train, y_train)

# 6. Evaluate and Save
y_pred = model.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)

# Save the trained model and vectorizer for future use
joblib.dump(model, 'sentiment_model.pkl')
joblib.dump(vectorizer, 'vectorizer.pkl')
"""
    add_code(doc, detailed_train_code.strip())

    # Prediction
    add_heading(doc, '5. src/core/predict.py (Making New Guesses)', level=1)
    add_paragraph(doc, "Now that the AI is smart and saved to a file, this script loads it up to analyze brand-new comments.")
    add_code(doc, "model, vectorizer = ModelManager().get_models()\n\ntext_features = vectorizer.transform([cleaned_text])\nprediction = model.predict(text_features)")
    doc.add_paragraph("• ModelManager: Ensures pre-trained model and vectorizer are loaded from disk or trained automatically.")
    doc.add_paragraph("• vectorizer.transform: Turns the new user comment into math numbers exactly the same way we did during training.")
    doc.add_paragraph("• model.predict: The AI looks at the numbers and spits out a final guess: Positive, Negative, or Neutral.")
    doc.add_paragraph("• CSV Export: Saves predictions into 'predicted_sentiments.csv' for downstream analytics.")

    # Statistics
    add_heading(doc, '6. src/analytics/statistics.py (Crunching the Numbers)', level=1)
    add_paragraph(doc, "Once we have our predictions, we aggregate the numbers.")
    add_code(doc, "counts = df[\"sentiment\"].value_counts()\npositive_pct = (positive_count / total) * 100")
    doc.add_paragraph("• We use Pandas to read the predicted dataset and count how many comments fell into each bucket.")
    doc.add_paragraph("• We calculate frequency counts and text length metrics for dashboard display.")

    # UI Manager
    add_heading(doc, '7. src/ui/ui_manager.py & app.py (The Desktop Dashboard)', level=1)
    add_paragraph(doc, "Users interact with a clean desktop dashboard built using Tkinter.")
    add_code(doc, "import tkinter as tk\nfrom src.ui.ui_manager import SentimentApp\n\nroot = tk.Tk()\napp = SentimentApp(root)\nroot.mainloop()")
    doc.add_paragraph("• SentimentApp: Renders the home screen, file upload controls, loading states, and analytics summary cards and tables.")
    doc.add_paragraph("• When the user clicks 'ANALYZE SENTIMENTS', our app runs the full pipeline in a background thread and updates the UI!")

    # Save Document
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.join(base_dir, "..", "..", "docs")
    os.makedirs(docs_dir, exist_ok=True)
    
    output_path = os.path.join(docs_dir, "Model_Explanation.docx")
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    main()
